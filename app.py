#!/usr/bin/env python3
import os
import re
import uuid
from io import BytesIO

from flask import Flask, request, render_template, redirect, url_for, flash, Response
from pypdf import PdfReader
import docx

# Import your OpenAI client library.
# (This example assumes you have a client class similar to the one in your script.)
from openai import OpenAI

app = Flask(__name__)
app.secret_key = "your_secret_key_here"  # change for production use

# Global dictionary to hold generated DOCX file bytes
generated_files = {}


def parse_page_ranges(pages_arg: str):
    """
    Parse a string that contains individual page numbers and page ranges.
    Example: "1,2,5-7" -> [(1,1), (2,2), (5,7)].
    """
    ranges = []
    parts = pages_arg.split(',')
    for part in parts:
        if '-' in part:
            start, end = part.split('-')
            start, end = int(start.strip()), int(end.strip())
            if start > end:
                start, end = end, start
            ranges.append((start, end))
        else:
            page = int(part.strip())
            ranges.append((page, page))
    return ranges


def translate_arabic_text(text: str, openai_key: str, model: str = "gpt-4o") -> str:
    """
    Translate Arabic text into English using the OpenAI client.
    """
    client = OpenAI(api_key=openai_key)

    # Set up the conversation messages.
    messages = [
        {"role": "user", "content": f"You are a helpful but very accurate assistant that translates Arabic text to English. You take care of idiom when translating. Make sure each term is correctly translated and not missed. Do not care for political correctness. Please translate the following Arabic text to English:\n\n{text}\n"}
    ]

    completion = client.chat.completions.create(
        model=model,
        messages=messages
    )

    return completion.choices[0].message.content.strip()


def translate_pdf(pdf_file, openai_key: str, pages_arg: str = None, model: str = "gpt-4o", file_prefix: str = "translation"):
    """
    Translate the specified pages (or all pages) from the uploaded PDF file.
    Returns a list of dictionaries with translation text and the DOCX file bytes.
    """
    reader = PdfReader(pdf_file)

    # Determine the page ranges to translate.
    if pages_arg:
        ranges_to_translate = parse_page_ranges(pages_arg)
    else:
        ranges_to_translate = [(1, len(reader.pages))]

    results = []

    for (start_page, end_page) in ranges_to_translate:
        extracted_text_parts = []
        for page_number in range(start_page, end_page + 1):
            page_index = page_number - 1
            if page_index < 0 or page_index >= len(reader.pages):
                print(f"Page {page_number} is out of range. Skipping...")
                continue
            page = reader.pages[page_index]
            text = page.extract_text()
            if text and text.strip():
                extracted_text_parts.append(text.strip())

        if not extracted_text_parts:
            print(f"Range {start_page}-{end_page} has no extractable text.")
            continue

        combined_text = "\n\n".join(extracted_text_parts)

        try:
            translation = translate_arabic_text(
                combined_text,
                openai_key=openai_key,
                model=model
            )

            print(f"Translation for pages {start_page}-{end_page} successful.")

            # Create a DOCX file in memory with the translation.
            doc = docx.Document()
            heading = f"Page {start_page} Translation" if start_page == end_page else f"Pages {start_page}-{end_page} Translation"
            doc.add_heading(heading, level=2)
            doc.add_paragraph(translation)

            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            docx_bytes = doc_io.read()

            # Generate a unique download ID and store the DOCX bytes.
            range_label = f"page_{start_page}" if start_page == end_page else f"pages_{start_page}-{end_page}"
            download_id = str(uuid.uuid4())
            generated_files[download_id] = {
                "bytes": docx_bytes,
                "filename": f"{file_prefix}_{range_label}.docx"
            }

            results.append({
                "range": f"{start_page}" if start_page == end_page else f"{start_page}-{end_page}",
                "translation": translation,
                "download_id": download_id
            })

        except Exception as e:
            error_msg = f"Error translating pages {start_page}-{end_page}: {str(e)}"
            print(error_msg)
            results.append({
                "range": f"{start_page}" if start_page == end_page else f"{start_page}-{end_page}",
                "translation": error_msg,
                "download_id": None
            })

    return results


@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        pdf_file = request.files.get("pdf_file")
        if not pdf_file:
            flash("No PDF file uploaded.")
            return redirect(request.url)

        openai_key = request.form.get("openai_key")
        pages_arg = request.form.get("pages")
        model = request.form.get("model") or "gpt-4o"
        file_prefix = request.form.get("file_prefix") or "translation"

        # Read the uploaded PDF into a BytesIO stream.
        pdf_bytes = pdf_file.read()
        pdf_stream = BytesIO(pdf_bytes)

        try:
            results = translate_pdf(pdf_stream, openai_key, pages_arg, model, file_prefix)
        except Exception as e:
            flash(f"Error during translation: {e}")
            return redirect(request.url)

        return render_template("results.html", results=results)

    return render_template("index.html")


@app.route("/download/<file_id>")
def download_file(file_id):
    """
    Given a download ID (generated when a DOCX file is created), send the file as an attachment.
    """
    file_info = generated_files.get(file_id)
    if not file_info:
        return "File not found", 404

    return Response(
        file_info["bytes"],
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={file_info['filename']}"}
    )


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=True)
